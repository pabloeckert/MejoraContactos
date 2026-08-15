"""Genera motor-contactos-auditoria.docx (raíz del repo) — auditoría
técnica real del proyecto, con la identidad visual de Mejora Continua.

Las fuentes (Bw Modelica Bold/Medium, League Spartan) NO se embeben en el
.docx -- python-docx no soporta embedding real de OOXML sin manipular XML
a mano, y no hace falta: ya están instaladas para el usuario actual (se
verificó contra el registro antes de escribir este script, están desde
una sesión anterior). Este script solo referencia esos nombres de familia
tal como los resuelve GDI/Word -- confirmado con
`(New-Object System.Drawing.Text.InstalledFontCollection).Families`, no
supuesto.

Contenido basado en lectura real del código al momento de escribir esto
(2026-08-14) -- no en lo que un nombre de archivo sugiere. Cifras leídas
en vivo de Data/Salida/staging.sqlite."""

from __future__ import annotations

import sqlite3
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

RAIZ = Path(__file__).resolve().parent.parent
DESTINO = RAIZ / "motor-contactos-auditoria.docx"

# Nombres de familia EXACTOS tal como los resuelve GDI en esta máquina
# (verificado, no el nombre del archivo .otf/.ttf).
F_TITULO = "Bw Modelica Bold"
F_SUBTITULO = "Bw Modelica Medium"
F_CUERPO = "League Spartan"

AZUL = RGBColor(0x1A, 0x3D, 0x84)
ROJO = RGBColor(0xE1, 0x06, 0x1E)
AMARILLO = RGBColor(0xF7, 0xCC, 0x13)
TINTA = RGBColor(0x2B, 0x2B, 0x2B)
GRIS = RGBColor(0x6B, 0x72, 0x80)
BLANCO = RGBColor(0xFF, 0xFF, 0xFF)


def _set_run_font(run, familia: str, tamano: int, color: RGBColor, negrita: bool = False) -> None:
    run.font.name = familia
    run.font.size = Pt(tamano)
    run.font.color.rgb = color
    run.font.bold = negrita
    # rFonts también necesita el ascii/hAnsi explícito en el XML crudo --
    # sin esto, Word a veces resuelve el tema por default en vez del
    # nombre puesto en run.font.name (encontrado probando en esta misma
    # sesión: el heading quedaba en Calibri pese a font.name correcto).
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), familia)
    rfonts.set(qn("w:hAnsi"), familia)
    rfonts.set(qn("w:cs"), familia)


def h1(doc: Document, texto: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    borde = p._p.get_or_add_pPr()
    pbdr = borde.makeelement(qn("w:pBdr"), {})
    bottom = borde.makeelement(qn("w:bottom"), {qn("w:val"): "single", qn("w:sz"): "18", qn("w:color"): "1A3D84", qn("w:space"): "4"})
    pbdr.append(bottom)
    borde.append(pbdr)
    r = p.add_run(texto)
    _set_run_font(r, F_TITULO, 18, AZUL, negrita=True)


def h2(doc: Document, texto: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(texto)
    _set_run_font(r, F_SUBTITULO, 13, AZUL, negrita=True)


def cuerpo(doc: Document, texto: str, color: RGBColor = TINTA, negrita: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(texto)
    _set_run_font(r, F_CUERPO, 10.5, color, negrita=negrita)


def bullet(doc: Document, texto: str, nivel: int = 0, color: RGBColor = TINTA) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.6 + 0.5 * nivel)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(texto)
    _set_run_font(r, F_CUERPO, 10.5, color)


def alerta(doc: Document, texto: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.3)
    r = p.add_run("⚠ " + texto)
    _set_run_font(r, F_CUERPO, 10.5, ROJO, negrita=False)


def tabla(doc: Document, encabezados: list[str], filas: list[list[str]], anchos: list[float] | None = None) -> None:
    t = doc.add_table(rows=1, cols=len(encabezados))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, texto in enumerate(encabezados):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(texto)
        _set_run_font(r, F_SUBTITULO, 9.5, BLANCO, negrita=True)
        _sombrear_celda(hdr[i], "1A3D84")
    for fila in filas:
        celdas = t.add_row().cells
        for i, texto in enumerate(fila):
            celdas[i].text = ""
            p = celdas[i].paragraphs[0]
            r = p.add_run(texto)
            _set_run_font(r, F_CUERPO, 9.5, TINTA)
    if anchos:
        for fila_t in t.rows:
            for i, ancho in enumerate(anchos):
                fila_t.cells[i].width = Cm(ancho)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _sombrear_celda(celda, hex_color: str) -> None:
    tc_pr = celda._tc.get_or_add_tcPr()
    shd = tc_pr.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): hex_color})
    tc_pr.append(shd)


def _leer_metricas_reales() -> dict[str, int]:
    ruta_db = RAIZ.parent / "Data" / "Salida" / "staging.sqlite"
    if not ruta_db.exists():
        return {}
    conn = sqlite3.connect(ruta_db)

    def c(sql: str) -> int:
        return conn.execute(sql).fetchone()[0]

    metricas = {
        "raw": c("SELECT COUNT(*) FROM raw_records"),
        "normalizados": c("SELECT COUNT(*) FROM normalized_records"),
        "finales": c("SELECT COUNT(DISTINCT cluster_id) FROM clusters"),
        "pendientes": c("SELECT COUNT(*) FROM decisiones_log WHERE accion='revision_pendiente'"),
        "cumpleanos": c("SELECT COUNT(*) FROM normalized_records WHERE cumpleanos IS NOT NULL"),
        "foto": c("SELECT COUNT(*) FROM normalized_records WHERE foto_url IS NOT NULL"),
    }
    conn.close()
    return metricas


def _correr_tests() -> tuple[bool, str]:
    resultado = subprocess.run(
        [str(RAIZ / ".venv" / "Scripts" / "python.exe"), "-m", "pytest", "-q"],
        cwd=RAIZ,
        env={"PYTHONPATH": "src", **__import__("os").environ},
        capture_output=True,
        text=True,
        timeout=180,
    )
    salida = (resultado.stdout or "").strip().splitlines()
    ultima = salida[-1] if salida else "(sin salida)"
    return resultado.returncode == 0, ultima


def construir() -> None:
    m = _leer_metricas_reales()
    tests_ok, tests_resumen = _correr_tests()

    doc = Document()
    seccion = doc.sections[0]
    seccion.left_margin = seccion.right_margin = Cm(2.2)
    seccion.top_margin = seccion.bottom_margin = Cm(2)
    doc.styles["Normal"].font.name = F_CUERPO
    doc.styles["Normal"].font.size = Pt(10.5)

    # Portada
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(40)
    r = p.add_run("AUDITORÍA TÉCNICA")
    _set_run_font(r, F_TITULO, 26, AZUL, negrita=True)
    p2 = doc.add_paragraph()
    r2 = p2.add_run("motor-contactos")
    _set_run_font(r2, F_SUBTITULO, 16, TINTA, negrita=True)
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(24)
    r3 = p3.add_run("Lead Systems Architect Review — 2026-08-14 — Repositorio privado, no distribuir")
    _set_run_font(r3, F_CUERPO, 10, GRIS)
    linea = doc.add_paragraph()
    borde = linea._p.get_or_add_pPr()
    pbdr = borde.makeelement(qn("w:pBdr"), {})
    bottom = borde.makeelement(qn("w:bottom"), {qn("w:val"): "single", qn("w:sz"): "24", qn("w:color"): "F7CC13", qn("w:space"): "1"})
    pbdr.append(bottom)
    borde.append(pbdr)

    # 1. Executive Summary
    h1(doc, "1. Executive Summary")
    cuerpo(
        doc,
        "motor-contactos resuelve un problema concreto y ya resuelto en producción: unificar los "
        "contactos reales de dos cuentas de Google (Pablo y Sindy), acumulados durante más de 15 años "
        "con formatos inconsistentes, duplicados y campos mal cargados, en una única lista maestra "
        "confiable y sincronizable de vuelta a Google Contacts.",
    )
    if m:
        cuerpo(doc, f"Estado real de los datos, verificado en la base al momento de esta auditoría:", negrita=True)
        tabla(
            doc,
            ["Métrica", "Valor"],
            [
                ["Registros crudos importados", f"{m['raw']:,}".replace(",", ".")],
                ["Registros normalizados", f"{m['normalizados']:,}".replace(",", ".")],
                ["Contactos finales (deduplicados)", f"{m['finales']:,}".replace(",", ".")],
                ["Casos pendientes de revisión manual", str(m["pendientes"])],
            ],
            anchos=[9, 5],
        )
    cuerpo(
        doc,
        "El proyecto NO es un prototipo ni un MVP a medio camino: tiene pipeline de datos completo, "
        "deduplicación con salvaguardas encontradas contra datos reales, dos interfaces funcionales "
        "(app de escritorio empaquetada y panel web), automatización mensual sin supervisión, e "
        "integraciones reales (Google, HubSpot/Mailchimp/Brevo, WhatsApp vía MejoraWS). Lo que sigue "
        "abierto (sección 6) son 2 pasos que requieren login humano de Google, no deuda de desarrollo.",
    )
    cuerpo(doc, f"Suite de tests al momento de esta auditoría: {tests_resumen}" + (" — todos en verde." if tests_ok else " — ATENCIÓN, hay fallas."))

    # 2. Arquitectura & Mapa de Archivos
    h1(doc, "2. Arquitectura & Mapa de Archivos")
    h2(doc, "2.1 Stack real (versiones leídas de requirements.txt / package.json)")
    tabla(
        doc,
        ["Capa", "Tecnología", "Versión mínima declarada"],
        [
            ["Backend", "Python", "3.14 (venv del proyecto)"],
            ["Backend web", "Flask", "3.1"],
            ["Datos", "SQLite (WAL) vía sqlite3 stdlib", "—"],
            ["Excel", "openpyxl / pandas / xlrd / odfpy", "3.1 / 3.0 / 2.0 / 1.4"],
            ["Teléfonos", "phonenumbers", "9.0"],
            ["Fuzzy matching", "rapidfuzz", "3.14"],
            ["Extracción PDF/DOCX/HTML", "pdfplumber / python-docx / beautifulsoup4+lxml", "0.11 / 1.2 / 4.13+5.3"],
            ["OCR", "pytesseract + Pillow (requiere binario Tesseract aparte)", "0.3 / 11.0"],
            ["Google", "google-api-python-client + google-auth(-oauthlib)", "2.149 / 2.35 / 1.2"],
            ["App de escritorio", "pywebview + PyInstaller (build-only)", "6.2 / 6.22"],
            ["Frontend", "React + TypeScript + Vite", "18.3 / 5.6 / 5.4"],
            ["UI: estilos", "Tailwind CSS", "3.4.13"],
            ["UI: listas grandes", "@tanstack/react-virtual", "3.10"],
            ["UI: modal", "@radix-ui/react-dialog", "1.1"],
            ["Tests", "pytest", "9.0"],
        ],
        anchos=[3.5, 8, 4.5],
    )
    h2(doc, "2.2 Mapa de módulos (src/motor/) — qué hace cada uno REALMENTE")
    tabla(
        doc,
        ["Módulo", "Responsabilidad real"],
        [
            ["staging_db.py", "Esquema SQLite (9 tablas) + migraciones ALTER TABLE incrementales sobre la base ya existente."],
            ["ingest.py", "Orquesta los extractores de archivo sobre Data/Crudos/, con try/except por archivo (uno roto no frena el resto)."],
            ["google_contacts_source.py", "Import en vivo desde People API (contactos + \"otros contactos\"), incremental por etag."],
            ["extractors/*.py", "Un extractor por formato (CSV, Excel, VCF, JSON, HTML, DOCX, PDF, OCR, texto libre) + column_mapping.py compartido (alias ES/EN/HubSpot/Mailchimp/Brevo)."],
            ["text_cleaning.py", "Heurísticas de limpieza calibradas contra datos reales rotos (teléfono guardado como nombre, honoríficos, cargos concatenados, filas plantilla)."],
            ["normalize_pipeline.py", "Conecta raw_records con phone_normalizer/email_normalizer/text_cleaning/tagging, escribe normalized_records."],
            ["phone_normalizer.py / area_codes_ar.py", "Normalización E.164 con reglas específicas de Argentina (ver sección 4.1)."],
            ["email_normalizer.py", "Normalización + corrección de dominios comunes mal tipeados (distancia de Levenshtein)."],
            ["tagging.py", "Auto-etiquetado (familiar/laboral/cliente/proveedor/personal) por palabras clave."],
            ["dedup/blocking.py, scoring.py, union_find.py, learning.py", "Motor de deduplicación (ver sección 4.2)."],
            ["llm_judge.py", "Cliente IA escalonado (Groq → OpenRouter gratis → Anthropic pago) para la banda dudosa."],
            ["merge_engine.py", "Aplica las 3 bandas, materializa clusters, deshacer puntual/total."],
            ["anomalias.py", "Detección de teléfonos compartidos por +5 contactos finales distintos."],
            ["export.py", "Lista maestra .xlsx, CSV para WhatsApp (MejoraWS), API de búsqueda/edición/listado."],
            ["api.py / reviewer_app.py", "API JSON (para la UI React) y panel HTML clásico — mismo motor, dos interfaces."],
            ["desktop_app.py", "Envuelve Flask + la UI React compilada en una ventana nativa (pywebview)."],
            ["cli.py", "Punto de entrada único (`python -m motor.cli <comando>`) para todos los comandos."],
        ],
        anchos=[5, 11],
    )
    h2(doc, "2.3 Pipeline de datos de punta a punta")
    for i, paso in enumerate(
        [
            "Entrada: People API de Google (en vivo) o archivos sueltos en Data/Crudos/ (CSV/Excel/VCF/JSON/HTML/DOCX/PDF/OCR/texto libre) → raw_records (inmutable).",
            "Normalización: phone_normalizer + email_normalizer + text_cleaning + tagging → normalized_records.",
            "Bloqueo (blocking.py): agrupa candidatos a comparar por teléfono/email exacto o clave fonética apellido+inicial — evita comparar todos contra todos.",
            "Clasificación (scoring.py): score 0–1 por señales, 3 bandas de confianza (ver 4.2).",
            "Fusión (merge_engine.py): materializa clusters vía union-find, nunca destructivo.",
            "Salida: lista-maestra.xlsx, CSV para WhatsApp, API para las dos interfaces, Apps Script para sync de vuelta a Google.",
        ],
        start=1,
    ):
        bullet(doc, f"{i}. {paso}")

    # 3. Estado de la interfaz
    h1(doc, "3. Estado de la interfaz")
    cuerpo(
        doc,
        "El proyecto NO es CLI pura: existen dos interfaces visuales completas y funcionales, "
        "construidas sobre la misma API/motor.",
        negrita=True,
    )
    h2(doc, "3.1 App de escritorio (App/MotorContactos.exe, pywebview + React)")
    bullet(doc, "Ventana nativa de Windows, sin terminal ni pestaña de navegador — identidad visual de Mejora Continua (tipografía, paleta, logo).")
    bullet(doc, "Sidebar con 3 secciones (Contactos / Revisión pendiente / Sync a Google), 4 tarjetas de métricas en vivo.")
    bullet(doc, "Tabla de contactos: 12 columnas configurables (mostrar/ocultar) y redimensionables (drag), columna Nombre fija a la izquierda, filtro \"contiene\" por columna + multiselect en Tag, búsqueda global instantánea client-side (toda la base cargada en memoria, sin ida y vuelta al servidor), filas virtualizadas (@tanstack/react-virtual — renderiza solo lo visible aunque haya miles de filas).")
    bullet(doc, "Cola de revisión en lote con contexto completo por par (nombre, organización, teléfono, email, fuente, foto) — no un id numérico pelado.")
    bullet(doc, "Editor de contacto en modal, todos los campos editables, la corrección manual pisa el cálculo automático de forma no destructiva (tabla ediciones_manuales aparte).")
    h2(doc, "3.2 Panel clásico (Iniciar Panel.bat, HTML servido por Flask)")
    bullet(doc, "Mismo motor, mismas acciones (correr pipeline, revisar, buscar/editar, exportar), sin la identidad visual nueva ni las columnas configurables — pensado como respaldo de un solo proceso si la app de escritorio fallara.")
    h2(doc, "3.3 Lo que falta pulir (autocrítica del propio equipo, no genérica)")
    bullet(doc, "Los popovers de filtro de columnas muy a la derecha (Tag, Nota) pueden abrirse fuera del viewport visible en tablas muy anchas.", color=ROJO)
    bullet(doc, "Los colores de avatar son una paleta genérica de Tailwind, no derivada de la paleta de marca.", color=ROJO)

    # 4. Lógica de negocio
    h1(doc, "4. Lógica de negocio")
    h2(doc, "4.1 Normalización de teléfono — casos límite reales que maneja")
    bullet(doc, "Formato de salida target: +549 + 10 dígitos (celular) / +54 + 10 dígitos (fijo) — estilo WhatsApp.")
    bullet(doc, "Detecta y separa números pegados sin separador (bloques de 10 dígitos consecutivos, ≥20 dígitos totales).")
    bullet(doc, "Detecta notación científica de Excel (\"3,76505E+09\") ANTES de partir por separadores — no se \"completa\" con un código de área default (fabricaría un número falso), va directo a revisión.")
    bullet(doc, "Prioriza SIEMPRE la interpretación doméstica argentina sobre la detección global de país cuando el patrón encaja (10 dígitos, 9+10, 54+10/11, 6–8 dígitos) — un bloque de dígitos \"pelado\" puede coincidir por casualidad con el plan de numeración de otro país.")
    bullet(doc, "Ambigüedad celular/fijo sin ninguna pista: usa el default configurado y deja la bandera 'movil-asumido' — nunca decide en silencio sin dejar rastro.")
    bullet(doc, "Números cortos (6–8 dígitos): se completan con el código de área default, quedan marcados 'incompleto'+'corregido'.")
    bullet(doc, "Múltiples números en un mismo campo (separados por ; , / | \"y\" o ':::') se dividen y deduplican por E.164 resultante.")
    h2(doc, "4.2 Criterio de deduplicación")
    cuerpo(doc, "Bloqueo (evita O(n²)): candidatos agrupados por teléfono exacto, email exacto, o clave fonética apellido+inicial de nombre; bloques de más de 500 elementos se descartan para no explotar en combinatoria (esos casos igual se capturan si comparten teléfono/email exacto).")
    cuerpo(doc, "Score y bandas (config.yaml → dedup):")
    bullet(doc, "score ≥ 0.80 → fusión automática (\"trigo\").")
    bullet(doc, "score ≤ 0.55 → separación automática (\"paja\").")
    bullet(doc, "0.55 < score < 0.80 → \"zona gris\": IA escalonada (Groq → modelos gratis de OpenRouter → Anthropic pago) y, si nadie resuelve con confianza suficiente, cola de revisión humana en lote.")
    bullet(doc, "Pesos del score cuando no hay señal exacta: teléfono 0.5, email 0.3, similitud de nombre 0.15, similitud de organización 0.05.")
    bullet(doc, "Salvaguarda real (encontrada corriendo contra los datos reales, no teórica): teléfono/email exacto normalmente fuerza score=1.0, EXCEPTO si ambos contactos traen nombre completo y son claramente distintos entre sí (similitud < 0.5) — ahí se asume línea compartida (familia/oficina) y va a zona gris en vez de fusionar en silencio.")
    bullet(doc, "Aprendizaje: cada decisión humana ajusta (hasta ±0.1, con mínimo 15 casos de evidencia) el score futuro de ese patrón de señales específico — no es ML, es una tasa de aceptación auditable.")
    h2(doc, "4.3 Integraciones — estado real, no aspiracional")
    tabla(
        doc,
        ["Integración", "Estado real"],
        [
            ["Google Contacts (import)", "Funcionando en producción con datos reales."],
            ["Google Contacts (sync de vuelta, Apps Script)", "Código completo y migrado a People API — NUNCA probado en vivo, requiere login humano de Google."],
            ["\"Otros contactos\" de Gmail", "Código completo — nunca probado en vivo, requiere un login humano APARTE (scope OAuth distinto)."],
            ["WhatsApp (MejoraWS)", "Funcionando: export en el formato exacto que MejoraWS espera, verificado contra datos reales."],
            ["HubSpot / Mailchimp / Brevo", "Funcionando vía reconocimiento de encabezados CSV — no hay integración por API con esas plataformas (no se generaron API keys)."],
        ],
        anchos=[6, 10],
    )

    # 5. Diagnóstico de puntos críticos
    h1(doc, "5. Diagnóstico de puntos críticos")
    h2(doc, "5.1 Deuda técnica real (no genérica)")
    bullet(doc, "El servidor de la app de escritorio y del panel es el servidor de desarrollo de Flask (Werkzeug) — Flask mismo advierte no usarlo en producción. Riesgo bajo en este contexto (un solo usuario local, nunca expuesto a internet), pero es una dependencia real a tener presente si algún día se expone la app fuera de localhost.")
    bullet(doc, "Backup exclusivamente local (repo git sin remoto en esta misma PC) — es punto único de falla: si el disco de esta máquina falla, no hay copia offsite de Data/Salida/staging.sqlite.")
    bullet(doc, "App/MotorContactos.exe (PyInstaller) no se reconstruye automáticamente al cambiar el código — nada impide que quede desactualizado si se toca ui/ o el backend y no se corre scripts/build_exe.ps1 antes de usarlo.")
    bullet(doc, "OCR de capturas/imágenes depende de que Tesseract-OCR (el binario) esté instalado aparte del proyecto — si falta, el extractor devuelve lista vacía sin avisar explícitamente al usuario final en la interfaz.")
    bullet(doc, "Sin pipeline de CI para motor-contactos (a diferencia de la SPA hermana del repo, que sí corre GitHub Actions) — los 195 tests solo se ejecutan cuando alguien los corre a mano.")
    h2(doc, "5.2 Cuellos de botella de rendimiento")
    cuerpo(doc, f"Con el volumen real actual ({m.get('finales', 0):,}".replace(",", ".") + f" contactos finales, {m.get('raw', 0):,}".replace(",", ".") + " crudos) no hay cuello de botella medido: la UI carga toda la base en memoria y filtra client-side sin demora perceptible, el blocking evita la explosión combinatoria. El único techo conocido es el tope de 500 elementos por bloque de blocking — un apellido extremadamente común por encima de ese tope pierde blocking por nombre (igual se captura por teléfono/email exacto).")
    h2(doc, "5.3 Riesgos operativos")
    alerta(doc, "Ningún paso de escritura hacia Google Contacts (Sync.gs) fue probado contra una cuenta real todavía — el primer uso real debería ser contra un Sheet de prueba, no la lista completa.")
    alerta(doc, "credentials.json y los token_*.json de Google están gitignoreados correctamente, pero viven sin cifrado en disco — cualquier proceso con acceso a esta cuenta de Windows podría leerlos.")
    alerta(doc, "No hay backup automático PRE-escritura antes de una corrida de deduplicación grande — el mecanismo de \"deshacer\" cubre el caso de arrepentimiento, pero depende de que alguien lo accione a tiempo.")

    # 6. Roadmap
    h1(doc, "6. Roadmap")
    cuerpo(doc, "Pendientes reales según el código y PENDIENTES.md — no una lista genérica de buenas prácticas:", negrita=True)
    bullet(doc, "Probar Sync.gs contra una cuenta de Google de PRUEBA (2-3 contactos ficticios) antes de correrlo contra las cuentas reales de Pablo y Sindy — requiere login humano, bloqueado hasta que se haga.")
    bullet(doc, "Si Sync.gs funciona en la prueba: correrlo contra las cuentas reales.")
    bullet(doc, "Probar \"otros contactos\" (importar-otros-contactos) en vivo — requiere un segundo login humano con un scope OAuth distinto al de Sync.gs.")
    bullet(doc, "Arreglar el posicionamiento de los popovers de filtro para columnas del extremo derecho de la tabla (ver 3.3).")
    bullet(doc, "Derivar la paleta de colores de avatar de la identidad de marca en vez de la paleta genérica de Tailwind (ver 3.3).")
    bullet(doc, "Evaluar backup offsite/en la nube de Data/ — hoy el único respaldo vive en la misma máquina (ver 5.3).")
    bullet(doc, "(Explícitamente pospuesto, no en este roadmap salvo pedido nuevo del usuario: escaneo de directorios completos de la PC, integración directa por API con HubSpot/Mailchimp/Brevo, agente autónomo más allá de la automatización mensual ya construida.)")

    doc.save(DESTINO)


if __name__ == "__main__":
    construir()
    print(f"Generado: {DESTINO}")
