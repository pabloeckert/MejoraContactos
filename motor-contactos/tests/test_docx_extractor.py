from docx import Document

from motor.extractors.docx_extractor import extraer_docx


def _crear_docx_con_tabla(path):
    doc = Document()
    tabla = doc.add_table(rows=3, cols=3)
    encabezados = ["Nombre", "Telefono", "Correo"]
    for i, texto in enumerate(encabezados):
        tabla.rows[0].cells[i].text = texto
    tabla.rows[1].cells[0].text = "Juan"
    tabla.rows[1].cells[1].text = "3743504517"
    tabla.rows[1].cells[2].text = "juan@gmail.com"
    tabla.rows[2].cells[0].text = "Maria"
    tabla.rows[2].cells[1].text = "3764368724"
    tabla.rows[2].cells[2].text = ""
    doc.save(str(path))


def test_extrae_filas_de_una_tabla_con_encabezado(tmp_path):
    path = tmp_path / "contactos.docx"
    _crear_docx_con_tabla(path)

    registros = extraer_docx(path)

    assert len(registros) == 2
    assert registros[0].campos["nombre"] == "Juan"
    assert registros[0].campos["telefono_1"] == "3743504517"
    assert registros[0].campos["email_1"] == "juan@gmail.com"
    assert "email_1" not in registros[1].campos


def test_docx_sin_tablas_no_produce_registros(tmp_path):
    from docx import Document as D

    path = tmp_path / "sin_tabla.docx"
    doc = D()
    doc.add_paragraph("Un documento sin ninguna tabla de contactos.")
    doc.save(str(path))

    assert extraer_docx(path) == []
